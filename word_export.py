from io import BytesIO
from docx import Document
from docx.shared import Pt


def create_word_document(title: str, body: str) -> bytes:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    doc.add_heading(title, level=1)
    for block in body.split("\n"):
        line = block.strip()
        if not line:
            continue
        if line[:2].isdigit() and "." in line[:4]:
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
